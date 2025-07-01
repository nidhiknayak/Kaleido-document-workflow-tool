import pandas as pd
import camelot
import pdfplumber
from docx import Document
from docx.table import Table as DocxTable
import json
import logging
from typing import List, Dict, Union, Optional
from pathlib import Path
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TableExtractor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']

    def extract_tables(self, file_path: str) -> Dict[str, Union[List[Dict], str]]:
        """Extract tables from supported file formats"""
        try:
            file_path = Path(file_path)
            file_extension = file_path.suffix.lower()

            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")

            logger.info(f"Processing file: {file_path.name}")

            if file_extension == '.pdf':
                return self._extract_from_pdf(str(file_path))
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(str(file_path))

        except Exception as e:
            logger.error(f"Error extracting tables: {str(e)}")
            return {
                "tables": [],
                "error": str(e),
                "file_name": file_path.name if 'file_path' in locals() else "unknown",
                "status": "failed"
            }

    def _extract_from_pdf(self, file_path: str) -> Dict[str, Union[List[Dict], str]]:
        """Extract tables from PDF using Camelot and pdfplumber as fallback"""
        tables_data = []

        try:
            logger.info("Attempting extraction with Camelot...")
            camelot_tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')

            if len(camelot_tables) > 0:
                for i, table in enumerate(camelot_tables):
                    df = table.df
                    table_dict = self._process_dataframe(df, f"camelot_table_{i}")
                    if table_dict:
                        tables_data.append(table_dict)

            if not tables_data:
                logger.info("Camelot failed, trying pdfplumber...")
                tables_data.extend(self._extract_with_pdfplumber(file_path))

        except Exception as e:
            logger.warning(f"Camelot extraction failed: {e}, trying pdfplumber...")
            tables_data.extend(self._extract_with_pdfplumber(file_path))

        return {
            "tables": tables_data,
            "file_name": Path(file_path).name,
            "status": "success" if tables_data else "no_tables_found",
            "extraction_method": "camelot" if any("camelot" in t.get("source", "") for t in tables_data) else "pdfplumber"
        }

    def _extract_with_pdfplumber(self, file_path: str) -> List[Dict]:
        """Extract tables using pdfplumber"""
        tables_data = []

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()

                    for table_num, table in enumerate(tables):
                        if table and len(table) > 1:
                            headers = table[0] if table[0] else [f"col_{i}" for i in range(len(table[0]))]
                            rows = table[1:] if len(table) > 1 else []
                            headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(headers)]

                            if rows:
                                df = pd.DataFrame(rows, columns=headers)
                                table_dict = self._process_dataframe(
                                    df,
                                    f"pdfplumber_page_{page_num}_table_{table_num}"
                                )
                                if table_dict:
                                    tables_data.append(table_dict)

        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")

        return tables_data

    def _extract_from_docx(self, file_path: str) -> Dict[str, Union[List[Dict], str]]:
        """Extract tables from DOCX files"""
        tables_data = []

        try:
            doc = Document(file_path)

            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = ' '.join([p.text.strip() for p in cell.paragraphs])
                        row_data.append(cell_text)
                    table_data.append(row_data)

                if table_data and len(table_data) > 1:
                    headers = table_data[0]
                    rows = table_data[1:]
                    headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(headers)]
                    df = pd.DataFrame(rows, columns=headers)
                    table_dict = self._process_dataframe(df, f"docx_table_{table_num}")
                    if table_dict:
                        tables_data.append(table_dict)

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {
                "tables": [],
                "error": str(e),
                "file_name": Path(file_path).name,
                "status": "failed"
            }

        return {
            "tables": tables_data,
            "file_name": Path(file_path).name,
            "status": "success" if tables_data else "no_tables_found",
            "extraction_method": "python-docx"
        }

    def _process_dataframe(self, df: pd.DataFrame, table_id: str) -> Optional[Dict]:
        """Process and clean DataFrame data"""
        try:
            # Remove completely empty rows and columns
            df = df.dropna(how='all').dropna(axis=1, how='all')

            if df.empty or len(df) == 0:
                return None

            # Clean column names
            df.columns = [str(col).strip() if col else f"col_{i}" for i, col in enumerate(df.columns)]

            # Handle duplicate column names
            cols = pd.Series(df.columns)
            for dup in cols[cols.duplicated()].unique():
                cols[cols[cols == dup].index.values.tolist()] = [
                    dup + f'_{i}' if i != 0 else dup for i in range(sum(cols == dup))
                ]
            df.columns = cols

            return {
                "table_id": table_id,
                "headers": df.columns.tolist(),
                "rows": df.values.tolist(),
                "shape": df.shape,
                "source": table_id.split('_')[0],
                "data": df.to_dict('records')
            }

        except Exception as e:
            logger.error(f"Error processing DataFrame: {e}")
            return None

    def export_to_csv(self, tables_data: List[Dict], output_dir: str = None) -> List[str]:
        """Export tables to CSV files"""
        if not output_dir:
            output_dir = tempfile.gettempdir()

        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)

        csv_files = []

        for table in tables_data:
            try:
                df = pd.DataFrame(table['data'])
                csv_path = output_dir / f"{table['table_id']}.csv"
                df.to_csv(csv_path, index=False)
                csv_files.append(str(csv_path))
                logger.info(f"Exported CSV: {csv_path}")
            except Exception as e:
                logger.error(f"Error exporting table {table['table_id']} to CSV: {e}")

        return csv_files

    def export_to_json(self, extraction_result: Dict, output_path: str = None) -> str:
        """Export extraction results to JSON file"""
        if not output_path:
            output_path = Path(tempfile.gettempdir()) / "extracted_tables.json"
        else:
            output_path = Path(output_path)

        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(extraction_result, f, indent=2, ensure_ascii=False)

            logger.info(f"Exported JSON: {output_path}")
            return str(output_path)

        except Exception as e:
            logger.error(f"Error exporting to JSON: {e}")
            raise